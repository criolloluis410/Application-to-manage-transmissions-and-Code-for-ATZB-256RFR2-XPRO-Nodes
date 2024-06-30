######################################################################################################################
# Decalracion de librerias ##########################################################################
from PyQt5.QtGui import QIntValidator
from PyQt5.QtWidgets import QApplication, QMainWindow, QDialog,QTableWidgetItem
from FrmAplicacionPrincipal import *
import sys
import serial
import time
import math
from PyQt5.QtCore import QTimer,Qt
import os
from tkinter import Tk, filedialog
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import csv
import pandas as pd
from tkinter import filedialog
from openpyxl import Workbook
from openpyxl.styles import Alignment

class VentanaPrinipal(QMainWindow):
    def __init__(self):
        super().__init__()
    ######################################################################################################################
    # Declaracion variables,componentes y eventos##########################################################################
        # Ejecucucion de la ventana principal de la alicacion
        self.primeraventana = Ui_MainWindow()
        self.primeraventana.setupUi(self)
        # Temporizadores y componentes para el manejo de las barras de progreso
        self.timerRecepcion = QTimer(self)
        self.timerRecepcion.timeout.connect(self.actualizarContadorRecepcion)
        self.timerDescarga = QTimer(self)
        self.timerDescarga.timeout.connect(self.actualizarContadorDescarga)
        self.primeraventana.progressBarRecepcion.setValue(0)
        self.primeraventana.progressBarDescarga.setValue(0)
        # Manejo de eventos de cada nuno de los componentes tipo boton de la aplicacion
        self.primeraventana.btnConectar.clicked.connect(self.ejecutarConexion)
        self.primeraventana.btnTrasmision.clicked.connect(self.EnviarOrdenes)
        self.primeraventana.btnConsulta.clicked.connect(self.ObtenerDatos)
        self.primeraventana.btnGenerarExel.clicked.connect(self.GenerarArchivoEXEL)
        self.primeraventana.btnGenerarWord.clicked.connect(self.GenerarArchivoWORD)
        self.primeraventana.btnBuscarCSV.clicked.connect(self.BuscarArchivoCSV)
        self.primeraventana.btnConvertirCSVaEXEL.clicked.connect(self.ConvertirArchivoCSVaEXEL)
        self.primeraventana.btnCancelarCVS.clicked.connect(self.CancelarCSVaEXEL)
        # Manejo de eventos de cada nuno de los componentes tipo radio_button de la aplicacion
        self.primeraventana.radioButton_1.toggled.connect(self.on_radio_button_toggled)
        self.primeraventana.radioButton_2.toggled.connect(self.on_radio_button_toggled)
        self.primeraventana.actionSalir.triggered.connect(self.Salir)
        #Funcion para mostrar la interfaz grafica
        self.show()
        # Variables y listas Globales
        self.puertoSerial = ''  # Puerto serial, almacena el puerto COM
        self.conexion = serial.Serial()  # Conexion serial, almacena la conexion serial previamente establecida
        self.DatosCompletosGlobales = b''# Variable para almacenar los datos obtenidos de la conexion serial
        self.contPruebas=0 # Variable para indicar el numero de veces que se le da ordenes al nodo transmisor para iniciar una prueba
        self.selected_ctr_cont_pr_value = 1 # Variable para manejar el contador de pruebas de manera manual o automatica
        self.parpadeoPB = 1 # Variable para mostrar un mensaje el cual indica que se esta realizando una transmicion o recepcion
        # Funciones Iniciales
        self.desaparecerComponentes()
        self.bloquearComponentesHojaDatos()
    ######################################################################################################################
    #Declaracion de funciones###########################################################################################
    #Funcion para cerrar la aplicacion
    def Salir(self):
        self.close()

    # Funcion para controlar el conteo de las pruebas para que sean manuales o automaticas
    def on_radio_button_toggled(self):
        if self.primeraventana.radioButton_1.isChecked()== True:
            self.selected_ctr_cont_pr_value = 1
            self.primeraventana.lineEditNumeroPrueba.setEnabled(True)
            print("Valor seleccionado:", self.selected_ctr_cont_pr_value)
            print("Contador Pruebas:", self.contPruebas)
        if self.primeraventana.radioButton_2.isChecked()== True:
            self.selected_ctr_cont_pr_value = 2
            self.primeraventana.lineEditNumeroPrueba.setEnabled(False)
            self.contPruebas=int(self.primeraventana.lineEditNumeroPrueba.text())
            self.primeraventana.label_NoPrueba.setText(self.primeraventana.lineEditNumeroPrueba.text())
            print("Valor seleccionado:", self.selected_ctr_cont_pr_value)
            print("Contador Pruebas:", self.contPruebas)

    # Funcion para bloquear componentes de la interfaz grafica de la aplicacion para evitar errores
    def desaparecerComponentes(self):
        self.primeraventana.btnTrasmision.setEnabled(False)
        self.primeraventana.progressBarRecepcion.setEnabled(False)
        self.primeraventana.spinBoxNumberTransmissions.setEnabled(False)
        self.primeraventana.spinBoxTimeBetweenFrames.setEnabled(False)
        self.primeraventana.spinBoxNumeroBloques.setEnabled(False)
        self.primeraventana.spinBoxTiempoEntreBloques.setEnabled(False)
        self.primeraventana.spinBoxPotenciaTransmision.setEnabled(False)
        self.primeraventana.btnConsulta.setEnabled(False)
        self.primeraventana.progressBarDescarga.setEnabled(False)
        self.primeraventana.btnGenerarWord.setEnabled(False)
        self.primeraventana.btnGenerarWord.setStyleSheet("color: rgb(255, 255, 255);\n""background - color: rgb(85, 170, 127);")
        self.primeraventana.label_Word.setStyleSheet("background-color: rgb(213, 227, 249);")
        self.primeraventana.lineEditNumeroPrueba.setValidator(QIntValidator())
        self.primeraventana.lineEditNumeroPrueba.setEnabled(False)
        self.primeraventana.radioButton_1.setEnabled(False)
        self.primeraventana.radioButton_2.setEnabled(False)
        self.primeraventana.btnGenerarExel.setEnabled(False)
        self.primeraventana.btnGenerarExel.setStyleSheet("color: rgb(255, 255, 255);\n""background - color: rgb(85, 170, 127);")
        self.primeraventana.label_Exel.setStyleSheet("background-color: rgb(170, 170, 127);")
        self.primeraventana.comboxTamanioTrama.setEnabled(False)
        self.primeraventana.label_DescargandoTramas.hide()
        self.primeraventana.label_EnviandoTramas.hide()
        self.primeraventana.btnCancelarCVS.hide()
        self.primeraventana.btnConvertirCSVaEXEL.hide()
        self.primeraventana.btnBuscarCSV.setEnabled(False)
        #background - color: rgb(85, 170, 127);

    # Funcion para desbloquear componentes de la interfaz grafica una vez la aplicacion se conecta con el nodo controlador
    def reaparecerComponentesPrimerEnvio(self):
        self.primeraventana.btnTrasmision.setEnabled(True)
        self.primeraventana.progressBarRecepcion.setEnabled(True)
        self.primeraventana.spinBoxNumberTransmissions.setEnabled(True)
        self.primeraventana.spinBoxTimeBetweenFrames.setEnabled(True)
        self.primeraventana.spinBoxNumeroBloques.setEnabled(True)
        self.primeraventana.spinBoxTiempoEntreBloques.setEnabled(True)
        self.primeraventana.spinBoxPotenciaTransmision.setEnabled(True)
        self.primeraventana.btnConsulta.setEnabled(True)
        self.primeraventana.progressBarDescarga.setEnabled(True)
        self.primeraventana.btnBuscarCSV.setEnabled(True)
        self.primeraventana.radioButton_1.setEnabled(True)
        self.primeraventana.radioButton_2.setEnabled(True)
        self.primeraventana.comboxTamanioTrama.setEnabled(True)
        self.primeraventana.lineEditNumeroPrueba.setEnabled(True)

    # Funcion para manejar la conexion de la aplicacion con el nodo controlador
    def ejecutarConexion(self):
        try:
            SeleccionDePuerto = self.primeraventana.comboxPorts.itemText(self.primeraventana.comboxPorts.currentIndex())
            self.puertoSerial = SeleccionDePuerto[0:3] + SeleccionDePuerto[4]
            # Configuracion de conexion serial
            self.conexion = serial.Serial(self.puertoSerial, 9600, 8, 'N', stopbits=1, timeout=None)
            time.sleep(1)
            consulta1 = b'\b'
            tramaHexString = ''
            codificado = bytes.fromhex(tramaHexString)  # tramaHexString.encode()
            self.conexion.write(codificado + consulta1)  # Señal para obtener datos
            print(f"MSG enviado: {codificado + consulta1}")
            dirr = self.conexion.read_until(b'\xFF')
            print('Direccion del nodo Conectado: ' + dirr[:2].hex())
            self.conexion.close()
            print('Conexión Exitosa')
            self.reaparecerComponentesPrimerEnvio()
        except:
            print('Conexión Fallida !')

    def EnviarOrdenes(self):
        try:# Configuracion de conexion serial nueva para evitar errores de conexion previos al envio de ordenes atraves del controlador
            SeleccionDePuerto = self.primeraventana.comboxPorts.itemText(self.primeraventana.comboxPorts.currentIndex())
            self.puertoSerial = SeleccionDePuerto[0:3] + SeleccionDePuerto[4]
            self.conexion = serial.Serial(self.puertoSerial, 9600, 8, 'N', stopbits=1, timeout=None)
            time.sleep(1)
            consulta2 = b'\b'
            tramaHexString = ''
            codificado = bytes.fromhex(tramaHexString)  # tramaHexString.encode()
            self.conexion.write(codificado + consulta2)  # Señal para obtener datos
            print(f"MSG enviado: {codificado + consulta2}")
            dirr = self.conexion.read_until(b'\xFF')
            print('Direccion del nodo Conectado: ' + dirr[:2].hex())
            #self.primeraventana.label_IDNodoConectado.setText("0x" + dirr[:2].hex())
            self.conexion.close()
            print('Conexión Exitosa')
        except:
            print('Conexión Fallida !')
        #Envio de tramas con cada uno de las caracteriticas que definiran el numero de tramas que debe transmitir el nodo Transmisor
        self.conexion.open()
        self.conexion.flush()
        sparador = b'\t'
        identificador = b'\n'
        mensaje = 'ini'  # Inicio de Ordenes
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Orden 'a' Numero de Transmiciones
        mensaje = self.primeraventana.spinBoxNumberTransmissions.text() + 'a'
        NumeroTX = int(self.primeraventana.spinBoxNumberTransmissions.text())
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Orden 'b' Tiempo entre tramas
        mensaje = self.primeraventana.spinBoxTimeBetweenFrames.text() + 'b'
        TiempoEntreTr = int(self.primeraventana.spinBoxTimeBetweenFrames.text())
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Orden 'c' Numero de Bloques
        mensaje = self.primeraventana.spinBoxNumeroBloques.text() + 'c'
        NumeroBl = int(self.primeraventana.spinBoxNumeroBloques.text())
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Orden 'd' Tiempo entre Bloques
        mensaje = self.primeraventana.spinBoxTiempoEntreBloques.text() + 'd'
        TiempoEntreBl = int(self.primeraventana.spinBoxTiempoEntreBloques.text())
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Orden 'e' Potencia de transmicion
        mensaje = self.primeraventana.spinBoxPotenciaTransmision.text() + 'e'
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Orden 'f' Potencia de transmicion
        tamañoTrama = self.primeraventana.comboxTamanioTrama.itemText(
            self.primeraventana.comboxTamanioTrama.currentIndex())
        ordenTamañoTrama = '0'
        if tamañoTrama == 'Short Frame':
            ordenTamañoTrama = '1'
        elif tamañoTrama == 'Long Frame':
            ordenTamañoTrama = '2'
        mensaje = ordenTamañoTrama + 'f'
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        # Finalizacion de Ordenes
        mensaje = 'fin'
        mensajeCodificado = mensaje.encode()
        self.conexion.write(mensajeCodificado + sparador)
        conf = self.conexion.read_until(b'\xFF')
        print('Confirmacion: ' + conf[:5].decode())
        self.conexion.close()  # Cerrar conexión
        print('Ordenes Completas ')
        #Configuracion de tiempos para calcular la duracion del envio del nodo tranmisor
        tiempoAdicionalBl = (TiempoEntreBl) * 0.70
        tiempoAdicionalTr = (TiempoEntreTr) * 0.70
        self.tiempoTotal=0.0
        if NumeroBl==1:
            self.tiempoTotal = (TiempoEntreTr*NumeroTX) / 1000
        elif NumeroBl>1:
            #self.tiempoTotal = ((TiempoEntreTr * NumeroTX) / 1000) + (TiempoEntreBl * NumeroBl)
            self.tiempoTotal = (((TiempoEntreTr+tiempoAdicionalTr) * (NumeroTX)* (NumeroBl)) / 1000) + ((TiempoEntreBl+ tiempoAdicionalBl) * (NumeroBl-1))
        print(f"tiempoTotal: {self.tiempoTotal}")
        self.iniciarTemporizadorRecepcion()

    # Funcion para iniciar el funcionamiento de la barra de progreso de Recepcion a traves de un temporizador
    def iniciarTemporizadorRecepcion(self):
        self.total_pasos = 100
        self.contador = 0
        tiempo_en_segundos = float(self.tiempoTotal)
        self.primeraventana.progressBarRecepcion.setValue(0)
        self.timer_intervalo = int(tiempo_en_segundos * 1000 / self.total_pasos)
        self.timerRecepcion.start(self.timer_intervalo)

    # Funcion para controlar la barra de progreso que indica la posible duracion del envio de tramas por parte del transmisor
    def actualizarContadorRecepcion(self):
        self.contador += 1
        self.primeraventana.progressBarRecepcion.setValue(self.contador)# Actualiza la barra de progreso

        self.primeraventana.label_EnviandoTramas.show()
        if self.parpadeoPB == 1:
            self.primeraventana.label_EnviandoTramas.setText("Nodo Transmisor Enviando Tramas.")
        elif self.parpadeoPB == 2:
            self.primeraventana.label_EnviandoTramas.setText("Nodo Transmisor Enviando Tramas..")
        elif self.parpadeoPB == 3:
            self.primeraventana.label_EnviandoTramas.setText("Nodo Transmisor Enviando Tramas...")

        if self.parpadeoPB == 1:
            self.parpadeoPB = 2
            #print("parpadeo ",self.parpadeoPB)
        elif self.parpadeoPB == 2:
            self.parpadeoPB = 3
            #print("parpadeo ", self.parpadeoPB)
        elif self.parpadeoPB == 3:
            self.parpadeoPB = 1
            #print("parpadeo ", self.parpadeoPB)
        # Condicional para detener el avanze de la barra de progreso lo cual indica el termino del envio de tramas
        if self.contador == self.total_pasos:
            self.timerRecepcion.stop()
            if self.selected_ctr_cont_pr_value == 1:
                self.primeraventana.label_NoPrueba.setText(self.primeraventana.lineEditNumeroPrueba.text())
            elif self.selected_ctr_cont_pr_value == 2:
                self.contPruebas=self.contPruebas+1
                self.primeraventana.lineEditNumeroPrueba.setText(str(self.contPruebas))
                #self.primeraventana.lineEditNumeroPrueba.setEnabled(False)
                self.primeraventana.label_NoPrueba.setText(str(self.contPruebas))
            self.primeraventana.btnGenerarWord.setEnabled(True)
            self.primeraventana.btnGenerarWord.setStyleSheet(
                "background-color: rgb(0, 41, 122);\n""color: rgb(255, 255, 255);\n""")
            self.primeraventana.label_Word.setStyleSheet("background-color: rgb(0, 67, 202);")
            self.parpadeoPB=1
            self.primeraventana.label_EnviandoTramas.hide()

    # Funcion para conocer si un nodo esta conectado a la aplicacion y conocer si este nodo puede enviar datos
    def ObtenerDatos(self):
        self.conexion.open()
        self.conexion.flush()
        consulta3 = b'\b'
        tramaHexString = ''
        codificado = bytes.fromhex(tramaHexString)  # tramaHexString.encode()
        self.conexion.write(codificado + consulta3)  # Señal para obtener datos
        print(f"MSG enviado: {codificado + consulta3}")
        dirr = self.conexion.read_until(b'\xFF')
        print('Direccion del nodo Conectado: ' + dirr[:2].hex())
        self.primeraventana.label_IDNodoConectado.setText("0x"+dirr[:2].hex())
        self.conexion.close()
        self.iniciarTemporizadorDescarga()

    # Funcion para iniciar el funcionamiento de la barra de progreso para indicar la descarga de datos a traves de un temporizador
    def iniciarTemporizadorDescarga(self):
        PruebaNo = self.primeraventana.lineEditNumeroPrueba.text()
        self.primeraventana.label_NoPrueba.setText(PruebaNo)
        self.conexion.open()
        self.conexion.flush()
        sparador = b'\r'
        tramaHexString = ''
        codificado = bytes.fromhex(tramaHexString)  # tramaHexString.encode()
        self.conexion.write(codificado + sparador)  # Señal para obtener datos
        print(f"MSG enviado: {codificado + sparador}")

        self.contador=0
        self.total_pasos = 40 # Establece las veces que se repetira el timer
        self.AvanzePB=0.0
        self.DatosCompletos = b''
        self.primeraventana.progressBarDescarga.setValue(0)
        self.timer_intervalo = 100 # 200 ms
        self.timerDescarga.start(self.timer_intervalo)

    # Funcion para controlar la barra de progreso que indica la recepcion de datos obtenidos de un nodo conectado a la app
    def actualizarContadorDescarga(self):
        #Manejo de los contadores para mostrar el avanze de la barra de progreso
        self.primeraventana.label_DescargandoTramas.show()
        self.contador += 1
        print(f"Contador: {self.contador}")
        DatosControlador = self.conexion.readline(255)
        print(f"DatosControlador: {DatosControlador}")
        self.DatosCompletos  = self.DatosCompletos  + DatosControlador
        self.primeraventana.progressBarDescarga.setValue(self.contador)
        valorAvanzePB = 100 / self.total_pasos
        self.AvanzePB = self.AvanzePB + valorAvanzePB
        self.primeraventana.progressBarDescarga.setValue(math.trunc(self.AvanzePB))
        if self.parpadeoPB == 1:
            self.primeraventana.label_DescargandoTramas.setText("Descargando Tramas.")
        elif self.parpadeoPB == 2:
            self.primeraventana.label_DescargandoTramas.setText("Descargando Tramas..")
        elif self.parpadeoPB == 3:
            self.primeraventana.label_DescargandoTramas.setText("Descargando Tramas...")

        if self.parpadeoPB == 1:
            self.parpadeoPB = 2
            #print("parpadeo ",self.parpadeoPB)
        elif self.parpadeoPB == 2:
            self.parpadeoPB = 3
            #print("parpadeo ", self.parpadeoPB)
        elif self.parpadeoPB == 3:
            self.parpadeoPB = 1
            #print("parpadeo ", self.parpadeoPB)
        # Condicional para detener el avanze de la barra de progreso lo cual indica el termino de la recepcion de datos
        if self.contador >= self.total_pasos:
            self.timerDescarga.stop()
            self.primeraventana.progressBarDescarga.setValue(100)
            self.conexion.close()
            self.primeraventana.btnGenerarExel.setEnabled(True)
            self.primeraventana.btnGenerarExel.setStyleSheet(
                "color: rgb(255, 255, 255);\n""background-color: rgb(0, 74, 11);")
            self.primeraventana.label_Exel.setStyleSheet("background-color: rgb(0, 74, 11);")
            print("DatosCompletos:")
            print(self.DatosCompletos )
            self.primeraventana.label_DescargandoTramas.hide()
            self.parpadeoPB = 1
            self.DatosCompletosGlobales=self.DatosCompletos

    # Funcion para agregar un borde de pagina a un archivo de word
    def add_page_border(self, document, border_width=6, border_color="000000", margin=0.5):
        sectPr = document.sections[0]._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), str(border_width))
        top.set(qn('w:color'), border_color)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(border_width))
        bottom.set(qn('w:color'), border_color)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(border_width))
        left.set(qn('w:color'), border_color)
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), str(border_width))
        right.set(qn('w:color'), border_color)
        pgBorders.append(top)
        pgBorders.append(bottom)
        pgBorders.append(left)
        pgBorders.append(right)
        sectPr.insert(0, pgBorders)
        # Adjust margins
        sections = document.sections
        for section in sections:
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)

    # Función para agregar texto centrado con un tamaño de letra específico
    def agregar_texto_centralizado(self, doc, texto, font_size=12):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(texto)
        run.bold = True
        font = run.font
        font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Función para agregar texto en negrita alineado a la izquierda
    def agregar_texto_negrita(self, doc, texto, font_size=12):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(texto)
        run.bold = True
        font = run.font
        font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Función para agregar una tabla con bordes y centrado en el contenido de las celdas (horizontal y vertical)
    def agregar_tabla(self, doc, datos):
        tabla = doc.add_table(rows=len(datos), cols=len(datos[0]))
        tabla.style = 'Table Grid'
        # Ajusta el ancho de las celdas al contenido solo si contiene números
        for i, row_data in enumerate(datos):
            row = tabla.rows[i].cells
            for j, text in enumerate(row_data):
                cell = row[j]
                if j % 2 != 0:
                    cell.width = Inches(0.1)  # Establece el ancho a 0 para que se ajuste al contenido
                cell.text = text
                # Centra horizontalmente el contenido de la celda
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Centra verticalmente el contenido de la celda
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        tabla.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Función para agregar una tabla con bordes y centrado en el contenido de las celdas (horizontal y vertical)
    def agregar_tabla_2(self, doc, datos):
        tabla = doc.add_table(rows=len(datos), cols=len(datos[0]))
        tabla.style = 'Table Grid'
        # Ajusta el ancho de las celdas al contenido solo si contiene números
        for i, row_data in enumerate(datos):
            row = tabla.rows[i].cells
            for j, text in enumerate(row_data):
                cell = row[j]
                if j == 0:
                    cell.width = Inches(0.1)  # Establece el ancho a 0 para que se ajuste al contenido
                cell.text = text
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centra el contenido

        tabla.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Función para crear un archivo word con todos los datos de transmision enviados al nodo transmisor
    def GenerarArchivoWORD(self):
        # Create a new Word document
        doc = Document()
        # Add border to the page and control the page border distance
        self.add_page_border(doc, margin=0.5)
        # Add centered text before New Title 1
        # Add a blank line
        doc.add_paragraph()
        self.agregar_texto_centralizado(doc, 'HOJA DE DATOS', font_size=16)
        # Add Title 1 and corresponding table
        self.agregar_texto_negrita(doc, '    Parámetros de transmisión', font_size=14)

        nTR=self.primeraventana.spinBoxNumberTransmissions.text()
        tTR=self.primeraventana.spinBoxTimeBetweenFrames.text()
        nBL=self.primeraventana.spinBoxNumeroBloques.text()
        tBL=self.primeraventana.spinBoxTiempoEntreBloques.text()
        pTX=self.primeraventana.spinBoxPotenciaTransmision.text()
        lTR=self.primeraventana.comboxTamanioTrama.itemText(self.primeraventana.comboxTamanioTrama.currentIndex())
        lTR=lTR[6:11]
        table_data_1 = [
            ['Numero de tramas por bloque:', nTR, 'Tiempo entre tramas[mseg]:', tTR],
            ['Numero de bloques:', nBL, 'Tiempo entre bloques[seg]:', tBL],
            ['Potencia de transmisiones[dbm]:', pTX, 'Tamaño de tramas:', lTR]
        ]
        column_widths_1 = [0.5, 0.5, 0.5, 0.5]  # Adjust column widths as needed
        self.agregar_tabla(doc, table_data_1)  # Adjust row height as needed
        # Add a blank line
        doc.add_paragraph()
        # Add Title 2 and corresponding table
        self.agregar_texto_negrita(doc, '    Localización de nodos', font_size=14)
        table_data_2 = [
            ['SRC_ADDR', 'Latitud', 'Longitud', 'Haltura'],
            ['0x0001', '0', '0', '0'],
            ['0x0002', '0', '0', '0'],
            ['0x0003', '0', '0', '0'],
            ['0x0004', '0', '0', '0'],
            ['0x0005', '0', '0', '0'],
            ['0x0006', '0', '0', '0'],
            ['0x0007', '0', '0', '0'],
            ['0x0008', '0', '0', '0']
        ]
        column_widths_2 = [1.5, 1, 1]  # Adjust column widths as needed
        self.agregar_tabla_2(doc, table_data_2)  # Adjust row height as needed
        # Add a blank line
        doc.add_paragraph()
        # Add notes
        self.agregar_texto_negrita(doc, '    Datos adicionales', font_size=14)
        notes_data = [
            ['Nodo_0x0001', ' '],
            ['Nodo_0x0002', ' '],
            ['Nodo_0x0003', ' '],
            ['Nodo_0x0004', ' '],
            ['Nodo_0x0005', ' '],
            ['Nodo_0x0006', ' '],
            ['Nodo_0x0007', ' '],
            ['Nodo_0x0008', ' ']
        ]
        notes_column_widths = [1.5, 1]  # Adjust column widths as needed
        self.agregar_tabla_2(doc, notes_data)  # Adjust row height as needed
        # Save the document
        root = Tk()
        root.withdraw()  # Hide the root window
        # Create directory 'PRUEBAS' if it doesn't exist
        directory = "PRUEBAS\Prueba No. "+self.primeraventana.lineEditNumeroPrueba.text()
        if not os.path.exists(directory):
            os.makedirs(directory)
        default_file_name = "Hoja de Datos No. "+self.primeraventana.lineEditNumeroPrueba.text()+".docx"   # Specify default file name
        file_path = filedialog.asksaveasfilename(initialdir=directory, initialfile=default_file_name,
                                                 defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if file_path:
            doc.save(file_path)
            print(
                f"Documento guardado exitosamente como '{file_path}' con bordes en cada página, contenido y texto centrado antes de Nuevo Título 1.")
            PruebaNo=self.primeraventana.lineEditNumeroPrueba.text()
            self.primeraventana.label_NoPrueba.setText(PruebaNo)
        else:
            print("Guardado cancelado.")

    # Función para crear un archivo de exel con todos los datos recibidos de cada nodo
    def GenerarArchivoEXEL(self):
        cadena_bytes = self.DatosCompletosGlobales
        # Inicializar la lista para almacenar los resultados
        resultados = []
        print(cadena_bytes)
        print(len(cadena_bytes))
        # Iterar sobre la cadena de bytes
        i = 0
        while i < len(cadena_bytes):

            # Buscar 'C0x' o 'L0x'
            if cadena_bytes[i:i + 3] == b'C0x':
                num_bytes_siguientes = 20
                try:
                    resultado = [
                        cadena_bytes[i:i + 1].decode(),  # L
                        cadena_bytes[i + 1:i + 7].decode(),  # 0x0000
                        float(cadena_bytes[i + 7:i + 11].decode()),  # 2.55
                        int(cadena_bytes[i + 11:i + 14].decode()),  # 003
                        ord(cadena_bytes[i + 14:i + 15]) - 90,  # Restar (-90)
                        cadena_bytes[i + 15:i + 23].decode()  # '1111111111' decodificado
                    ]
                finally:
                    # Manejo de la excepción
                    print("Datos incompletos")
                resultados.append(resultado)
                i += 3 + num_bytes_siguientes
            elif cadena_bytes[i:i + 3] == b'L0x':
                num_bytes_siguientes = 97
                resultado = [
                    cadena_bytes[i:i + 1].decode(),  # L
                    cadena_bytes[i + 1:i + 7].decode(),  # 0x0000
                    float(cadena_bytes[i + 7:i + 11].decode()),  # 2.55
                    int(cadena_bytes[i + 11:i + 14].decode()),  # 003
                    ord(cadena_bytes[i + 14:i + 15]) - 90,  # Restar (-90)
                    cadena_bytes[i + 15:i + 100].decode()  # '1111111111' decodificado
                ]
                resultados.append(resultado)
                i += 3 + num_bytes_siguientes
            else:
                i += 1

        contTramas=0
        print('Datos encontrados:')
        for resultado in resultados:
            print(resultado)
            contTramas+=1
        # Crear un DataFrame con los resultados
        self.primeraventana.label_NoTramasRecibidas.setText(str(contTramas))
        PruebaNo = self.primeraventana.lineEditNumeroPrueba.text()
        self.primeraventana.label_NoPrueba.setText(PruebaNo)

        df = pd.DataFrame(resultados,
                          columns=['Tipo', 'No. Secuencia RX', 'Nivel Bateria[V]', 'Ptx[dBm]', 'Prx[dBm]', 'Carga Util'])
        print(df)
        # Crear una ventana de Tkinter (no se mostrará físicamente)
        # Crear una ventana de diálogo para seleccionar la ubicación y el nombre del archivo Excel
        root = Tk()
        root.withdraw()  # Ocultar la ventana principal de la aplicación

        numTx = self.primeraventana.label_NoPrueba.text()
        directory = "PRUEBAS\Prueba No. " + numTx
        if not os.path.exists(directory):
            os.makedirs(directory)
        idNodo=self.primeraventana.label_IDNodoConectado.text()

        default_file_name = "Tramas descargadas del nodo " + idNodo + ".xlsx"  # Specify default file name
        #file_path = filedialog.asksaveasfilename(initialdir=directory, initialfile=default_file_name, defaultextension=".docx", filetypes=[("Word Document", "*.docx")])

        file_path = filedialog.asksaveasfilename(initialdir=directory, initialfile=default_file_name, defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        #print('paths:'+file_path)
        print('paths:')
        # Verificar si se proporcionó un nombre de archivo antes de guardar
        if file_path:
            # Guardar el DataFrame en el archivo Excel seleccionado
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Hoja1')

                # Ajustar automáticamente el tamaño de las celdas
                worksheet = writer.sheets['Hoja1']
                for column_cells in worksheet.columns:
                    max_length = 0
                    column = [cell for cell in column_cells]
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = (max_length + 2)
                    worksheet.column_dimensions[column[0].column_letter].width = adjusted_width

                # Centrar el contenido en todas las celdas
                for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row, min_col=1,
                                               max_col=worksheet.max_column):
                    for cell in row:
                        cell.alignment = cell.alignment.copy(horizontal='center', vertical='center')

            print(f"DataFrame guardado exitosamente en {file_path}")
            self.primeraventana.btnGenerarExel.setEnabled(False)
            self.primeraventana.btnGenerarExel.setStyleSheet(
                "color: rgb(255, 255, 255);\n""background - color: rgb(85, 170, 127);")
            self.primeraventana.label_Exel.setStyleSheet("background-color: rgb(170, 170, 127);")
        else:
            print("Operación de guardado cancelada.")

    # Función para bucar el archivo CSV generado por el nodo transmisor
    def BuscarArchivoCSV(self):
        ruta_archivo = filedialog.askopenfilename(filetypes=[("Archivos CSV", "*.csv")])
        if ruta_archivo:
            print("Ruta del archivo seleccionado:", ruta_archivo)
            self.rutaArchivoCSV=ruta_archivo
            self.primeraventana.btnCancelarCVS.show()
            self.primeraventana.btnConvertirCSVaEXEL.show()
            self.primeraventana.btnBuscarCSV.hide()
        else:
            print("Ruta Incorrecta")

    # Función para leer cada una de las filas de los archivos CSV encontrados
    def leer_csv(self, nombre_archivo):
        datos_csv = []
        with open(nombre_archivo, 'r') as archivo_csv:
            lector_csv = csv.reader(archivo_csv, delimiter=',')
            for fila in lector_csv:
                datos_csv.append(fila)
        return datos_csv

    # Función para convertir un archivo CSV a un EXEL
    def convertir_a_excel(self, datos_csv):
        numeroSecuencia_ascii = []
        latitud = []
        longitud = []
        print(datos_csv)
        for fila in datos_csv:
            print(fila)
            if 'Trama ASCII' in fila[0]:
                trama = fila[0].split(': ')[1]
                trama = "0x"+trama
                print("trama: ", trama)
                numeroSecuencia_ascii.append(trama)

            if 'Datos GPS' in fila[0]:
                latitud_csv = fila[0].split(': ')
                longitud_csv = fila[1].split(': ')
                latitud.append(latitud_csv[2])
                longitud.append(longitud_csv[1])

        df = pd.DataFrame({
            'No. Secuencia ARD-ATZB': numeroSecuencia_ascii,
            'Latitud': latitud,
            'Longitud': longitud
        })

        numTx2 = self.primeraventana.label_NoPrueba.text()
        directory = "PRUEBAS\Prueba No. " + numTx2
        if not os.path.exists(directory):
            os.makedirs(directory)

        default_file_name = "Datos del Nodo Transmisor(GPS) Prueba No. " + numTx2+ ".xlsx"  # Specify default file name

        file_path = filedialog.asksaveasfilename(initialdir=directory, initialfile=default_file_name, defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        #print('paths:'+file_path)
        # Verificar si se proporcionó un nombre de archivo antes de guardar
        if file_path:
            # Guardar el DataFrame en el archivo Excel seleccionado

            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False)
                # Obtener la hoja de cálculo creada
                workbook = writer.book
                worksheet = writer.sheets['Sheet1']
                # Ajustar el tamaño de las celdas al contenido
                for column in worksheet.columns:
                    max_length = 0
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = (max_length + 2) * 1.2
                    worksheet.column_dimensions[column[0].column_letter].width = adjusted_width
                    # Centrar el contenido en las celdas
                    for row in worksheet.iter_rows(min_row=2, min_col=1, max_row=worksheet.max_row,
                                                   max_col=worksheet.max_column):
                        for cell in row:
                            cell.alignment = Alignment(horizontal='center', vertical='center')
            print("Archivo Excel a partir de CSV generado con éxito:", default_file_name)

    # Función para controlar los componentes de la apliacion los cuales permiten convertir los archivos CSV a EXEL
    def ConvertirArchivoCSVaEXEL(self):
        datos_csv = self.leer_csv(self.rutaArchivoCSV)
        self.convertir_a_excel(datos_csv)
        self.primeraventana.btnCancelarCVS.hide()
        self.primeraventana.btnConvertirCSVaEXEL.hide()
        self.primeraventana.btnBuscarCSV.show()

    # Función para cancelar la convercion de archivos CSV
    def CancelarCSVaEXEL(self):
        self.primeraventana.btnCancelarCVS.hide()
        self.primeraventana.btnConvertirCSVaEXEL.hide()
        self.primeraventana.btnBuscarCSV.show()
        self.rutaArchivoCSV=""

    # Funcion para bloquear componentes para uso a futuro de la aplicacion
    def bloquearComponentesHojaDatos(self):
        listaCHD=[]
        listaCHD.append(self.primeraventana.lineEdit_1)
        listaCHD.append(self.primeraventana.lineEdit_2)
        listaCHD.append(self.primeraventana.lineEdit_3)
        listaCHD.append(self.primeraventana.lineEdit_4)
        self.primeraventana.groupBoxLocalizacionNodos.setStyleSheet("color: rgb(120, 120, 120);\n")
        self.primeraventana.groupBoxLocalizacionNodos.setEnabled(False)
        for oblinrEd in listaCHD:
            oblinrEd.setEnabled(False)
        self.primeraventana.groupBoxDatosAdicionales.setStyleSheet("color: rgb(120, 120, 120);\n")
        self.primeraventana.groupBoxDatosAdicionales.setEnabled(False)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ventanaP=VentanaPrinipal()
    ventanaP.show()
    sys.exit(app.exec_())
